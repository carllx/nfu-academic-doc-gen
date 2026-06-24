import docx

doc = docx.Document('05_Assessment_Generator/Template_Exam_Criteria_AB.docx')
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            # We want to remove empty paragraphs
            for p in list(cell.paragraphs):
                # if the paragraph has no text after stripping
                if not p.text.strip():
                    # remove it from the xml
                    p._element.getparent().remove(p._element)

doc.save('05_Assessment_Generator/Template_Exam_Criteria_AB.docx')
print("Done cleanup")
