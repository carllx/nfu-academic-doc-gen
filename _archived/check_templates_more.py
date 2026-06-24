import docx

def check_docx(filename):
    print(f"Checking {filename}...")
    try:
        doc = docx.Document(filename)
        for p in doc.paragraphs:
            if '【' in p.text or '_' in p.text:
                print(f"  Para: {p.text}")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '【' in cell.text or '_' in cell.text:
                        print(f"  Cell: {cell.text.strip()}")
    except Exception as e:
        pass

check_docx("05_Assessment_Generator/Template_Exam_Checklist_A.docx")
check_docx("05_Assessment_Generator/Template_Exam_Criteria_A.docx")
