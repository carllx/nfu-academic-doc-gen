import docx
import sys

def check_docx(filename):
    print(f"Checking {filename}...")
    try:
        doc = docx.Document(filename)
        found = False
        for p in doc.paragraphs:
            if '{' in p.text or '<' in p.text or '[' in p.text:
                print(f"  Para: {p.text}")
                found = True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{' in cell.text or '<' in cell.text or '[' in cell.text:
                        # Print only non-empty and unique?
                        print(f"  Cell: {cell.text.strip()}")
                        found = True
        if not found:
            print("  No obvious variables found.")
    except Exception as e:
        print(f"  Error reading {filename}: {e}")

check_docx("05_Assessment_Generator/Template_Exam_Checklist_A.docx")
check_docx("05_Assessment_Generator/Template_Exam_Criteria_A.docx")
check_docx("05_Assessment_Generator/Template_Exam_Checklist_B.docx")
check_docx("05_Assessment_Generator/Template_Exam_Criteria_B.docx")
