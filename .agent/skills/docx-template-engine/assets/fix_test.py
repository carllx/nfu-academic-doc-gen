from docx import Document
import json

doc = Document()
doc.add_paragraph('Hello {{name}}')
doc.add_paragraph('This is null: {{null_text}}')
doc.add_paragraph('This is zero: {{zero_val}}')
doc.add_paragraph('This is false: {{false_val}}')
doc.add_paragraph('{% for item in table_items %}')
doc.add_paragraph('Item: {{item}}')
doc.add_paragraph('{% endfor %}')
doc.save('/Users/yamlam/Downloads/carllx-skills/docx-template-filler/assets/example_template.docx')

data = {
  "name": "Antigravity",
  "null_text": "",
  "zero_val": 0,
  "false_val": False,
  "table_items": ["apple", "banana"]
}
with open('/Users/yamlam/Downloads/carllx-skills/docx-template-filler/assets/example_data.json', 'w') as f:
    json.dump(data, f, indent=2)
