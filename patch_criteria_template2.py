import docx

doc = docx.Document('/Users/yamlam/Downloads/教务材料/05_Assessment_Generator/Template_Exam_Criteria_AB.docx')

for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if '{{ a_theme }}' in p.text:
                    if len(p.runs) > 0:
                        # Clear existing runs
                        text = p.text.replace('{{ a_theme }}', '一、创作主题：\n{{ a_theme }}')
                        for run in p.runs:
                            run.text = ''
                        p.runs[0].text = text
                if '{{ b_theme }}' in p.text:
                    if len(p.runs) > 0:
                        text = p.text.replace('{{ b_theme }}', '一、创作主题：\n{{ b_theme }}')
                        for run in p.runs:
                            run.text = ''
                        p.runs[0].text = text

doc.save('/Users/yamlam/Downloads/教务材料/05_Assessment_Generator/Template_Exam_Criteria_AB.docx')
print("Patched tables in Template_Exam_Criteria_AB.docx")
