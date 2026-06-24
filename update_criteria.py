import docx

doc = docx.Document('05_Assessment_Generator/Template_Exam_Criteria_AB.docx')
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            # We want to clear out the text "一、创作主题：" and "{{ a_theme }}" and "{{ b_theme }}"
            # It might be spread across multiple runs or paragraphs
            for p in cell.paragraphs:
                if '一、创作主题：' in p.text or '{{ a_theme }}' in p.text or '{{ b_theme }}' in p.text:
                    # To be safe, clear the paragraph
                    for r in p.runs:
                        r.text = r.text.replace('一、创作主题：', '').replace('{{ a_theme }}', '').replace('{{ b_theme }}', '')
                        
                # Wait, if they are in different runs, string matching in run might fail. 
                # Let's rebuild the paragraph text if we find these substrings in p.text.
                full_text = p.text
                if any(x in full_text for x in ['一、创作主题：', '{{ a_theme }}', '{{ b_theme }}', '{% if a_theme or a_imgs %}', '{% if b_theme or b_imgs %}']):
                    # Replace in the full text
                    new_text = full_text.replace('一、创作主题：', '')
                    new_text = new_text.replace('{{ a_theme }}', '')
                    new_text = new_text.replace('{{ b_theme }}', '')
                    new_text = new_text.replace('{% if a_theme or a_imgs %}', '{% if a_imgs %}')
                    new_text = new_text.replace('{% if b_theme or b_imgs %}', '{% if b_imgs %}')
                    
                    if new_text != full_text:
                        # Clear all runs and set the first run to new_text
                        if p.runs:
                            p.runs[0].text = new_text
                            for r in p.runs[1:]:
                                r.text = ''

doc.save('05_Assessment_Generator/Template_Exam_Criteria_AB.docx')
print("Done")
