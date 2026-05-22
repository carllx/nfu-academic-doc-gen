#!/usr/bin/env python3
"""
课程目标汇总文档生成器
将三门课程的课程目标整合到一个 Word 文档中
根据人培版本自动适配表格列数（3列/4列）
"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ═══ 课程数据 ═══
COURSES = [
    {
        "name": "交互产品综合创作",
        "semester": "2025-2026-1",
        "grade": "2022级",
        "major": "数字媒体艺术（专升本）",
        "plan": "2024版",
        "schema": "coarse",
        "columns": 3,
        "objectives": [
            ("知识目标", [
                ("1、掌握Figma的核心功能，包括自动布局、约束、组件与变体、交互原型制作等。", "毕业要求4.基本技能"),
                ("2、掌握Spline的基础操作，包括3D建模、材质灯光、状态事件、物理引擎和变量动画。", "毕业要求4.基本技能"),
                ("3、理解交互设计背后的心理学原理，如格式塔原则、情感化设计及主流交互设计定律。", "毕业要求4.基本技能\n毕业要求6.创新设计能力"),
                ("4、了解Apple HIG、Google Material Design等行业主流UI/UX设计语言的特点与应用。", "毕业要求4.基本技能\n毕业要求6.创新设计能力"),
                ("5、识记交互产品从概念到发布的基本流程与方法。", "毕业要求4.基本技能"),
            ]),
            ("能力目标", [
                ("1、能够应用 Figma 独立完成符合行业规范的 2D 网页高保真设计，并掌握基础的 HTML/CSS 结构认知，确保设计稿具备可开发性。", "毕业要求8.团队合作能力"),
                ("2、能够应用Spline创作出具有丰富交互性和视觉吸引力的3D网页作品。", "毕业要求6.创新设计能力"),
                ("3、具备分析并应用交互设计理论与心理学原则的能力，以优化产品易用性和用户体验。", "毕业要求6.创新设计能力"),
                ("4、具备综合运用 2D UI 设计与 3D 交互技术的能力，并能通过工具（如 Spline Export 或 基础前端代码）实现 Web 端的最终技术落地，输出符合考核标准的文件结构。", "毕业要求6.创新设计能力"),
                ("5、具备独立完成产品调研、策划、设计迭代及项目展示的综合项目实践能力。", "毕业要求8.团队合作能力"),
            ]),
            ("素质目标", [
                ("1.培养学生树立以用户为中心的设计思维，具备通过设计解决实际问题的意识。", "毕业要求6.创新设计能力"),
                ("2.培养学生具备批判性思维与创新精神，能够探索新颖的交互体验形式，并对新兴技术（如AI）有批判性的理解。", "毕业要求6.创新设计能力"),
                ("3.培养学生具备团队协作精神与良好的沟通能力，适应现代设计工作流程。", "毕业要求8.团队合作能力"),
                ("4.培养学生树立设计师的职业道德和社会责任感，理解并践行包容性设计与无障碍设计原则，抵制\u201c暗黑模式\u201d等不道德设计，以科技向善为准则，创作有益于社会发展的数字内容。", "毕业要求6.创新设计能力\n毕业要求8.团队合作能力"),
            ]),
        ],
    },
    {
        "name": "三维动画基础-数字建模",
        "semester": "2025-2026-1",
        "grade": "2023级",
        "major": "数字媒体艺术（专升本）",
        "plan": "2025版",
        "schema": "detailed",
        "columns": 4,
        "objectives": [
            ("知识目标", [
                ("1.熟练掌握Maya软件的核心功能模块及其操作规范。", "毕业要求2.专业知识", "观测点2.1：掌握数字媒体艺术基础知识，熟练使用三维建模相关硬件和软件工具，精通建模和动画等核心技能。"),
                ("2.深入理解三维建模的基础原理与关键技术要点。", "毕业要求2.专业知识", "观测点2.1：深入理解数字媒体艺术发展脉络及基础概念，为建模相关职业技能奠定坚实基础。"),
                ("3.精通材质系统的应用方法及其视觉表现技巧；", "毕业要求2.专业知识", "观测点2.3：具备全流程创作能力，能够配置材质以实现逼真的视觉效果，达到行业标准。"),
                ("4.全面了解主流渲染器的工作原理与参数优化设置。", "毕业要求2.专业知识", "观测点2.1：熟练掌握渲染工具及技术流程，理解渲染相关参数优化以满足行业需求。"),
            ]),
            ("能力目标", [
                ("1.掌握基础模型创建能力，熟练运用多边形与NURBS建模核心技术，能够高效构建并优化几何体。", "毕业要求4.实际应用能力", "观测点4.1：精通三维建模软件与技术，应用建模技能完成优化几何体的项目实践。"),
                ("2.精通材质表现技能，根据项目需求精准配置材质，实现模型质感的真实呈现，显著提升模型的真实感与视觉表现力。", "毕业要求4.实际应用能力", "观测点4.1：结合艺术与设计理论，应用材质和纹理设计，创作符合行业标准的高质量作品。"),
                ("3. 具备基础渲染能力，深入理解光影原理与渲染流程，能够完成高质量的场景渲染输出。", "毕业要求4.实际应用能力", "观测点4.2：深入掌握渲染技术，融合传统与新技术，提升渲染效率与作品表现力。"),
                ("4．培养项目实践能力，具备独立完成简单三维作品创作的能力，形成适应行业需求的基本应用能力。", "毕业要求4.实际应用能力", "观测点4.3：识别并解决项目创作中的问题，确保通过专业技能交付高质量成果。"),
            ]),
            ("素质目标", [
                ("1．艺术表现素养：培养三维造型的审美能力，熟练掌握传统艺术元素在现代三维设计中的创新应用。", "毕业要求3：创造性思维", "观测点3.1：融合艺术、科学与专业知识，创作具有审美吸引力的数字作品。"),
                ("2．问题解决能力：构建系统化的分析与解决问题的思维框架，形成高效、切实可行的技术解决方案。", "毕业要求4.实际应用能力", "观测点4.3：诊断并解决建模与渲染中的技术与创意问题，确保高效完成任务。"),
                ("3．职业发展能力：养成持续学习技术的习惯，建立规范化的工作流程，为职业发展奠定坚实基础。", "毕业要求9：学习与发展能力", "观测点9.1：树立终身学习意识，主动更新技能以适应数字媒体艺术行业的快速发展。"),
                ("4．严谨的工作习惯：在建模与渲染过程中，注重细节处理，秉持严谨的工作态度，确保作品的精确性与一致性。", "毕业要求1：品德素养", "观测点1.2：展现责任感与职业素养，通过严谨的工作流程确保高质量作品输出。"),
                ("5．文化传承能力：精通传统建筑元素的数字化表达方法，创作具有鲜明文化特色的三维场景，助力文化传承。", "毕业要求8：行业洞察力", "观测点8.1：汲取多元文化精华，结合数字媒体艺术特质，创作具有文化共鸣的作品，满足社会与市场需求。"),
            ]),
        ],
    },
    {
        "name": "交互影像创作",
        "semester": "2024-2025-2",
        "grade": "2022级",
        "major": "数字媒体艺术（专升本）",
        "plan": "2024版",
        "schema": "coarse",
        "columns": 3,
        "objectives": [
            ("知识目标", [
                ("1、理解交互影像创作的核心概念与原理，包括身体与技术对话的媒介、传感器技术应用、交互艺术对社会议题的批判性表达等内容；", "毕业要求4.基本技能"),
                ("2、掌握TouchDesigner的核心操作与理论，包括节点系统、数据流逻辑、参数绑定（表达式/导出模式）、六大元件家族（TOP/CHOP/SOP/DAT等）的功能与应用场景；", "毕业要求4.基本技能"),
                ("3、熟悉实时视觉效果开发流程，涵盖抽象影像生成（噪声渐变、粒子系统）、三维渲染（材质调控、灯光属性）、GLSL着色器编程等技术；", "毕业要求4.基本技能\n毕业要求6.创新设计能力"),
                ("4、了解音画互动技术与硬件整合方法，包括音频信号分析（滤波/调制）、Arduino串口通信、Kinect体感交互、DMX灯光协议等；", "毕业要求4.基本技能\n毕业要求5.综合应用能力"),
                ("5、掌握社会批判性思维在艺术创作中的融入，关注技术伦理、身体规训、集体行为变革等议题。", "毕业要求6.创新设计能力"),
            ]),
            ("能力目标", [
                ("1、能够使用TouchDesigner独立开发动态视觉内容（如抽象壁纸、数据驱动动画）、实时交互装置（体感控制、穿戴设备交互）及三维场景构建；", "毕业要求4.基本技能\n毕业要求5.综合应用能力"),
                ("2、具备硬件与软件整合能力，通过传感器（距离/光敏/运动）、Arduino控制器、Kinect等设备实现物理交互与数据输入；", "毕业要求4.基本技能\n毕业要求5.综合应用能力"),
                ("3、熟练运用网络通信协议（OSC/TCP/IP）实现多终端协同，支持远程控制与分布式艺术展演；", "毕业要求4.基本技能\n毕业要求5.综合应用能力"),
                ("4、掌握音画同步设计技巧，包括音频实时可视化、AI辅助音乐生成与视听表达的技术转化；", "毕业要求4.基本技能\n毕业要求6.创新设计能力"),
                ("5、完成从概念设计到落地的完整交互项目开发，涵盖技术方案制定、硬件部署、空间媒体化改造及效果调试；", "毕业要求5.综合应用能力\n毕业要求6.创新设计能力"),
                ("6、在团队协作中担任多元角色（编程、装置搭建、创意设计），适应跨学科项目协作流程。", "毕业要求7.沟通表达\n毕业要求8.团队合作能力"),
            ]),
            ("素质目标", [
                ("\u201c技术-身体-社会\u201d三位一体的创作思维，探索人机共生语境下的艺术表达可能性；", "毕业要求6.创新设计能力"),
                ("2、深化批判性思考能力，反思技术在公共空间中的规训作用，通过作品传递社会观察；", "毕业要求6.创新设计能力"),
                ("3、强化开放性与协作意识，通过集体创作模式构建多主体参与的艺术系统；", "毕业要求7.沟通表达\n毕业要求8.团队合作能力"),
                ("4、树立技术伦理观念，在硬件改造、数据采集与应用中遵循安全规范与隐私保护原则；", "毕业要求6.创新设计能力"),
                ("5、提升数字美学素养，平衡抽象影像的感官冲击与深层语义传达。", "毕业要求6.创新设计能力"),
            ]),
        ],
    },
]


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shd)


def add_table_3col(doc, objectives):
    """构建3列表格（coarse模式）"""
    # 计算总行数
    total_rows = 1  # 表头
    for dim, items in objectives:
        total_rows += len(items)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(10.0)
        row.cells[2].width = Cm(4.0)

    # 表头
    headers = ["课程目标", "课程目标", "对应毕业要求"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "黑体"
        run.font.size = Pt(10.5)
        set_cell_shading(cell, "D9E2F3")

    # 数据行
    for dim, items in objectives:
        start_row = len(table.rows)
        for idx, item in enumerate(items):
            row = table.add_row()
            # 维度名 - 只在第一行写
            if idx == 0:
                cell0 = row.cells[0]
                cell0.text = ""
                p = cell0.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(dim)
                run.bold = True
                run.font.name = "仿宋"
                run.font.size = Pt(10.5)

            # 目标内容
            cell1 = row.cells[1]
            cell1.text = ""
            p = cell1.paragraphs[0]
            run = p.add_run(item[0])
            run.font.name = "仿宋"
            run.font.size = Pt(10.5)

            # 毕业要求
            cell2 = row.cells[2]
            cell2.text = ""
            p = cell2.paragraphs[0]
            run = p.add_run(item[1])
            run.font.name = "仿宋"
            run.font.size = Pt(10.5)

        # 合并维度名单元格
        if len(items) > 1:
            a = table.cell(start_row, 0)
            b = table.cell(start_row + len(items) - 1, 0)
            a.merge(b)

    return table


def add_table_4col(doc, objectives):
    """构建4列表格（detailed模式）"""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(2.0)
        row.cells[1].width = Cm(6.0)
        row.cells[2].width = Cm(3.5)
        row.cells[3].width = Cm(5.0)

    # 表头
    headers = ["课程目标", "课程目标", "对应毕业要求", "对应观测点"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "黑体"
        run.font.size = Pt(10.5)
        set_cell_shading(cell, "D9E2F3")

    # 数据行
    for dim, items in objectives:
        start_row = len(table.rows)
        for idx, item in enumerate(items):
            row = table.add_row()
            # 维度名
            if idx == 0:
                cell0 = row.cells[0]
                cell0.text = ""
                p = cell0.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(dim)
                run.bold = True
                run.font.name = "仿宋"
                run.font.size = Pt(10.5)

            # 目标内容
            cell1 = row.cells[1]
            cell1.text = ""
            p = cell1.paragraphs[0]
            run = p.add_run(item[0])
            run.font.name = "仿宋"
            run.font.size = Pt(10.5)

            # 毕业要求
            cell2 = row.cells[2]
            cell2.text = ""
            p = cell2.paragraphs[0]
            run = p.add_run(item[1])
            run.font.name = "仿宋"
            run.font.size = Pt(10.5)

            # 观测点
            cell3 = row.cells[3]
            cell3.text = ""
            p = cell3.paragraphs[0]
            run = p.add_run(item[2] if len(item) > 2 else "")
            run.font.name = "仿宋"
            run.font.size = Pt(10.5)

        # 合并维度名单元格
        if len(items) > 1:
            a = table.cell(start_row, 0)
            b = table.cell(start_row + len(items) - 1, 0)
            a.merge(b)

    return table


def add_paragraphs_none(doc, objectives):
    """构建纯文字段落（none模式，2023版及以前）"""
    dim_labels = {"知识目标": "（一）知识目标", "能力目标": "（二）能力目标", "素质目标": "（三）素质目标"}
    for dim, items in objectives:
        # 分类小标题
        p = doc.add_paragraph()
        run = p.add_run(dim_labels.get(dim, dim))
        run.bold = True
        run.font.name = "仿宋"
        run.font.size = Pt(12)
        # 每条目标作为独立段落
        for item in items:
            desc = item[0] if isinstance(item, (list, tuple)) else item
            p = doc.add_paragraph()
            run = p.add_run(desc)
            run.font.name = "仿宋"
            run.font.size = Pt(12)


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = "仿宋"
    style.font.size = Pt(12)

    # 标题
    title = doc.add_heading("课程目标汇总", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 说明
    p = doc.add_paragraph()
    run = p.add_run("本文档汇总了三门课程的课程目标，根据各课程适用的人才培养方案版本，自动适配格式：\n"
                     "• 2023版及以前（none模式）→ 纯文字段落\n"
                     "• 2024版人培（coarse模式）→ 3列表格（维度名、课程目标、对应毕业要求）\n"
                     "• 2025版人培（detailed模式）→ 4列表格（维度名、课程目标、对应毕业要求、对应观测点）\n"
                     "专升本映射规则：专升本年级+2=对标普本人培版本")
    run.font.name = "仿宋"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    from datetime import datetime
    p = doc.add_paragraph()
    run = p.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.name = "仿宋"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 遍历课程
    for i, course in enumerate(COURSES):
        if i > 0:
            doc.add_page_break()

        # 课程标题
        h = doc.add_heading(f"{i+1}. {course['name']}", level=1)

        # 元信息
        meta_lines = [
            f"开课年级：{course['grade']}　　开课专业：{course['major']}",
            f"开课学期：{course['semester']}　　适用人培：{course['plan']}",
            f"渲染模式：{course['schema']}" + (f"（{course['columns']}列表格）" if course.get('columns') else "（纯段落）"),
        ]
        for line in meta_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "仿宋"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph()  # 空行

        # 课程目标小标题
        h2 = doc.add_heading("二、课程目标", level=2)

        # 根据 schema_type 选择输出模式
        schema = course.get("schema", "detailed")
        if schema == "none":
            add_paragraphs_none(doc, course["objectives"])
        elif schema == "coarse":
            add_table_3col(doc, course["objectives"])
        else:
            add_table_4col(doc, course["objectives"])

    # 保存
    out_path = "/Users/yamlam/Documents/nfu - 教务/课程目标汇总_2025-2026.docx"
    doc.save(out_path)
    size_kb = os.path.getsize(out_path) / 1024
    print(f"✅ 文档已生成: {out_path}")
    print(f"   文件大小: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
