from docx import Document
from pathlib import Path

def rebuild_template():
    base_path = Path('/Users/yamlam/Downloads/教务材料/05_Assessment_Generator')
    template_path = base_path / 'Template_Exam_Criteria_AB.docx'
    
    if not template_path.exists():
        print("Template not found!")
        return

    doc = Document(template_path)
    
    # 1. Update Table
    table = doc.tables[0]
    
    # Remove all rows except the header (row 0) and the first data row (row 1)
    for i in range(len(table.rows) - 1, 1, -1):
        row = table.rows[i]
        table._tbl.remove(row._tr)
        
    # Update row 1 with Jinja tags
    row = table.rows[1]
    row.cells[0].text = '{% tr for sec in sections %}{{ sec.content }}'
    row.cells[1].text = '{{ sec.total_score }}{% tr endfor %}'
    
    # 2. Update Paragraphs (Sections)
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if '一、' in p.text:
            start_idx = i
        if '参考示例' in p.text and start_idx != -1 and end_idx == -1:
            end_idx = i
            
    if start_idx != -1 and end_idx != -1:
        # Delete existing section paragraphs
        for i in range(end_idx - 1, start_idx - 1, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            
        ref_p = doc.paragraphs[start_idx] # This is now the "参考示例" paragraph
        
        # Insert Jinja loop for sections
        ref_p.insert_paragraph_before('{% for sec in sections %}')
        p_title = ref_p.insert_paragraph_before('{{ loop.index }}、{{ sec.section_name }}（{{ sec.total_score }}分）')
        p_title.runs[0].bold = True
        ref_p.insert_paragraph_before('{{ sec.content }}')
        ref_p.insert_paragraph_before('{% endfor %}')
        ref_p.insert_paragraph_before('')

    # 3. Update A/B Theme
    ab_p_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if 'A卷参考示例' in p.text or 'B卷参考示例' in p.text:
            ab_p_idx = i
            break
            
    if ab_p_idx != -1:
        p = doc.paragraphs[ab_p_idx]
        p._element.getparent().remove(p._element)
        
        doc.add_paragraph('{% if a_theme or a_img %}')
        doc.add_heading('A卷参考示例', level=2)
        doc.add_paragraph('{{ a_theme }}')
        doc.add_paragraph('{{ a_img }}')
        doc.add_paragraph('{% endif %}')
        
        doc.add_paragraph('{% if b_theme or b_img %}')
        doc.add_heading('B卷参考示例', level=2)
        doc.add_paragraph('{{ b_theme }}')
        doc.add_paragraph('{{ b_img }}')
        doc.add_paragraph('{% endif %}')

    # 4. Replace Course Name
    for p in doc.paragraphs:
        if '课程名称' in p.text:
            p.text = p.text.replace('课程名称', '课程名称：{{ course_name }}')
            
    doc.save(template_path)
    print("Successfully rebuilt Template_Exam_Criteria_AB.docx")

if __name__ == '__main__':
    rebuild_template()
